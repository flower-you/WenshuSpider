#!/usr/bin/env python
# -*- coding:utf8 -*-
# @TIME     :2019/5/18 13:39
# @Author   :17976
# @File     :piplines.py
# @Description:
import os

import pymongo
from docx import Document
from pymongo.errors import DuplicateKeyError

import settings

#存储到excel
class csv_write():
    def __init__(self,file_name='wenshu.csv'):
        self.file_name = file_name

    def write_item(self,item):
        with open(self.file_name, 'a+', encoding='utf-8-sig') as f:
            f.write(str(item['casedocid'])+',')
            f.write(str(item['casename']).replace("\t"," ")+',')
            f.write(str(item['casetype']).replace("\t", " ")+',')
            f.write(str(item['casejudgedate']).replace("\t", " ")+',')
            f.write(str(item['caseprocedure']).replace("\t", " ")+',')
            f.write(str(item['casenumber']).replace("\t", " ")+',')
            f.write(str(item['casenopublicreason']).replace("\t", " ")+',')
            f.write(str(item['casecontenttype']).replace("\t", " ")+',')
            f.write(str(item['caseuploaddate']).replace("\t", " ")+',')
            f.write(str(item['caseclosemethod']).replace("\t", " ")+',')
            f.write(str(item['caseeffectivelevel']).replace("\t", " ")+',')

            f.write(str(item['casecourt']['casecourtid']).replace("\t", " ")+',')
            f.write(str(item['casecourt']['casecourtname']).replace("\t", " ")+',')
            f.write(str(item['casecourt']['casecourtprovince']).replace("\t", " ")+',')
            f.write(str(item['casecourt']['casecourtcity']).replace("\t", " ")+',')
            f.write(str(item['casecourt']['casecourtdistrict']).replace("\t", " ")+',')
            f.write(str(item['casecourt']['casecourtarea']).replace("\t", " ")+',')

            f.write(str(item['casecontent']['casebasecontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['caseaddcontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['caseheadcontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['casemaincontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['casecorrectionscontent']).replace("\t"," ")+',')
            f.write(str(item['casecontent']['casedoccontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['caselitigationcontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['casepartycontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['casetailcontent']).replace("\t", " ")+',')
            f.write(str(item['casecontent']['caseresultcontent']).replace("\t", " ")+',')

            f.write(str(item['casecontent']['casestrcontent']).replace("\t", " ")+',')

            f.write('\n')


#存储到word
class DocStore():
    def __init__(self, file_name='wenshu.docx'):
        self.file_name = file_name
        self.document = Document()
        # self.document = Document(file_name)
        # 不存在会报错，存在也必须是上次通过python-dox成功写入的非空白文件，否则还是报错，可能格式有问题

    def write_item(self,item):
        self.document.add_heading(str(item['casename']),level=1)
        p1 = self.document.add_paragraph('')
        p1.add_run('文书id：').bold = True
        p1.add_run(str(item['casedocid'])).italic = True
        self.document.add_heading('案件信息', level=2)
        self.document.add_paragraph('  案件类型：'+str(item['casetype']))
        self.document.add_paragraph('  裁判日期：' + str(item['casejudgedate']))
        self.document.add_paragraph('  审判程序：' + str(item['caseprocedure']))
        self.document.add_paragraph('  案号：'+str(item['casenumber']))
        self.document.add_paragraph('  不公开理由：'+str(item['casenopublicreason']))
        self.document.add_paragraph('  文书全文类型：'+str(item['casecontenttype']))
        self.document.add_paragraph('  上传日期：'+str(item['caseuploaddate']))
        self.document.add_paragraph('  结案方式：'+str(str(item['caseclosemethod'])))
        self.document.add_paragraph('  效力层级：'+str(item['caseeffectivelevel']))

        self.document.add_heading('法院信息',level=2)
        self.document.add_paragraph('  法院ID：'+str(item['casecourt']['casecourtid']))
        self.document.add_paragraph('  法院名称：' + str(item['casecourt']['casecourtname']))
        self.document.add_paragraph('  法院省份：' + str(item['casecourt']['casecourtprovince']))
        self.document.add_paragraph('  法院地市：' + str(item['casecourt']['casecourtcity']))
        self.document.add_paragraph('  法院区县：' + str(item['casecourt']['casecourtdistrict']))
        self.document.add_paragraph('  法院区域：' + str(item['casecourt']['casecourtarea']))

        self.document.add_heading('内容信息', level=2)
        self.document.add_paragraph('  案件基本情况原文：'+str(item['casecontent']['casebasecontent']))
        self.document.add_paragraph('  附加原文：'+str(item['casecontent']['caseaddcontent']))
        self.document.add_paragraph('  文件首部段落原文：'+str(item['casecontent']['caseheadcontent']))
        self.document.add_paragraph(' 裁判要旨段原文：'+str(item['casecontent']['casemaincontent']))
        self.document.add_paragraph('  补全正文：'+str(item['casecontent']['casecorrectionscontent']))
        self.document.add_paragraph('  案件内容：'+str(item['casecontent']['casedoccontent']))
        self.document.add_paragraph('  诉讼记录段原文'+str(item['casecontent']['caselitigationcontent']))
        self.document.add_paragraph(' 诉讼参与人信息部分原文'+str(item['casecontent']['casepartycontent']))
        self.document.add_paragraph('  文件尾部原文：'+str(item['casecontent']['casetailcontent']))
        self.document.add_paragraph('  判决结果段原文：'+str(item['casecontent']['caseresultcontent']))

        self.document.add_heading('案件原文', level=2)
        self.document.add_paragraph('    ' + str(item['casecontent']['casestrcontent']))

    def complete(self):
        self.document.add_page_break()
        self.document.save(self.file_name)

#存储到mongoDb
class mongoStore():
    def __init__(self,table_name):
        host = settings.mongo_msg["MONGODB_HOST"]
        port = int(settings.mongo_msg["MONGODB_PORT"])
        dbname = settings.mongo_msg["MONGODB_DBNAME"]
        table = table_name
        #创建数据库连接
        self.client = pymongo.MongoClient(host=host, port=port)
        # 指定数据库
        mydb = self.client[dbname]
        #  # 设置文书ID为唯一索引,避免插入重复数据
        mydb[table].ensure_index('casedocid', unique=True)
        self.post = mydb[table]


    def query_docid(self,docid):
        return self.post.find_one({"casedocid":docid})

    def find_all(self):
        return self.post.find()


    def process_item(self, item):
       try:
           data = dict(item)
           self.post.insert(data)
           return item
       except DuplicateKeyError:
           # 索引相同,即为重复数据,捕获错误
           print('Duplicate key error collection')
           return item

    def close_client(self):
        self.client.close()




